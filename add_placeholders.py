# import os
# from docx import Document

# def replace_text_in_doc(doc, old_text, new_text):
#     for paragraph in doc.paragraphs:
#         if old_text in paragraph.text:
#             inline = paragraph.runs
#             for i in range(len(inline)):
#                 if old_text in inline[i].text:
#                     text = inline[i].text.replace(old_text, new_text)
#                     inline[i].text = text

#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     if old_text in paragraph.text:
#                         inline = paragraph.runs
#                         for i in range(len(inline)):
#                             if old_text in inline[i].text:
#                                 text = inline[i].text.replace(old_text, new_text)
#                                 inline[i].text = text

# def add_placeholders_to_docx(folder_path):
#     placeholders = {
#         "{{ PROGRAM_NAME }}": "",
#         "{{ CREATED_BY }}": "",
#         "{{ CHANGE_NUMBER }}": "",
#         "{{ JOB_LOG_NUMBER }}": "",
#         "{{ TECHNICAL_NAME }}": "",
#         "{{ DESCRIPTION }}": "",
#         "{{ TEST_CONDITION }}": "",
#         "{{ CUSTOMER_REQUIREMENT }}": "",
#         "{{ TEST_PLAN_PREPARED_BY }}": "",
#         "{{ TEST_PLAN_REVIEWED_BY }}": "",
#         "{{ TESTING_BY }}": "",
#         "{{ TESTING_REVIEWED_BY }}": "",
#         "{{ TEST_RESULT_PREPARED_BY }}": "",
#         "{{ TEST_RESULT_REVIEWED_BY }}": "",
#         "{{ SCREENSHOT_OUTPUT }}": "",
#         "07.07.2025": "{{ DATE }}"
#     }

#     for filename in os.listdir(folder_path):
#         if filename.endswith(".docx"):
#             filepath = os.path.join(folder_path, filename)
#             doc = Document(filepath)
#             for old_text, new_text in placeholders.items():
#                 replace_text_in_doc(doc, old_text, new_text)
#             doc.save(filepath)
#             print(f"Processed {filename}")

# if __name__ == "__main__":
#     folder_to_process = "charm-docs-template"
#     add_placeholders_to_docx(folder_to_process)

import os
from docx import Document

def replace_text_in_doc(doc, old_text, new_text):
    """Replaces all occurrences of old_text with new_text in a document."""
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            inline = paragraph.runs
            # Replace strings and retain formatting
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(old_text, new_text)
                    inline[i].text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        inline = paragraph.runs
                        # Replace strings and retain formatting
                        for i in range(len(inline)):
                            if old_text in inline[i].text:
                                text = inline[i].text.replace(old_text, new_text)
                                inline[i].text = text

def add_placeholders_to_docx(folder_path):
    """Adds placeholders to the Test Plan and Test Results documents."""
    
    # Define replacements for Test Plan.docx
    test_plan_replacements = {
        "/ 07.07.2025": "/ {{DATE}}",
        "Jyosyula Siva Amrutha": "{{TEST_PLAN_REVIEWED_BY}}",
        "Swagata Roy": "{{TESTING_BY}}",
        "Indraneel Mazumder": "{{TESTING_REVIEWED_BY}}"
    }

    # Define replacements for Test Results.docx
    test_results_replacements = {
        ": / 07.07.2025": ": / {{DATE}}",
        "Swagata Roy": "{{TEST_RESULT_REVIEWED_BY}}",
        "Indraneel Mazumder": "{{TESTING_REVIEWED_BY}}",
        "TESTING BY: Swagata Roy": "TESTING BY: {{TESTING_BY}}" # More specific replacement
    }

    # Process Test Plan.docx
    test_plan_path = os.path.join(folder_path, "Test Plan.docx")
    if os.path.exists(test_plan_path):
        print("Processing Test Plan.docx...")
        doc = Document(test_plan_path)
        for old, new in test_plan_replacements.items():
            replace_text_in_doc(doc, old, new)
        doc.save(test_plan_path)
        print("Finished processing Test Plan.docx.")
    else:
        print("Warning: Test Plan.docx not found.")

    # Process Test Results.docx
    test_results_path = os.path.join(folder_path, "Test Results.docx")
    if os.path.exists(test_results_path):
        print("Processing Test Results.docx...")
        doc = Document(test_results_path)
        for old, new in test_results_replacements.items():
            replace_text_in_doc(doc, old, new)
        doc.save(test_results_path)
        print("Finished processing Test Results.docx.")
    else:
        print("Warning: Test Results.docx not found.")


if __name__ == "__main__":
    folder_to_process = "charm-docs-template"
    add_placeholders_to_docx(folder_to_process)