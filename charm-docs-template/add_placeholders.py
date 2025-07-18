
import os
from docx import Document

def replace_text_in_doc(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(old_text, new_text)
                    inline[i].text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        inline = paragraph.runs
                        for i in range(len(inline)):
                            if old_text in inline[i].text:
                                text = inline[i].text.replace(old_text, new_text)
                                inline[i].text = text

def add_placeholders_to_docx(folder_path):
    placeholders = {
        "{{ PROGRAM_NAME }}": "",
        "{{ CREATED_BY }}": "",
        "{{ CHANGE_NUMBER }}": "",
        "{{ JOB_LOG_NUMBER }}": "",
        "{{ TECHNICAL_NAME }}": "",
        "{{ DESCRIPTION }}": "",
        "{{ TEST_CONDITION }}": "",
        "{{ CUSTOMER_REQUIREMENT }}": "",
        "{{ TEST_PLAN_PREPARED_BY }}": "",
        "{{ TEST_PLAN_REVIEWED_BY }}": "",
        "{{ TESTING_BY }}": "",
        "{{ TESTING_REVIEWED_BY }}": "",
        "{{ TEST_RESULT_PREPARED_BY }}": "",
        "{{ TEST_RESULT_REVIEWED_BY }}": "",
        "{{ SCREENSHOT_OUTPUT }}": "",
        "07.07.2025": "{{ DATE }}"
    }

    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            filepath = os.path.join(folder_path, filename)
            doc = Document(filepath)
            for old_text, new_text in placeholders.items():
                replace_text_in_doc(doc, old_text, new_text)
            doc.save(filepath)
            print(f"Processed {filename}")

if __name__ == "__main__":
    folder_to_process = "charm-docs-template"
    add_placeholders_to_docx(folder_to_process)
