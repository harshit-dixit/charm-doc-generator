import pandas as pd
from docx import Document
from docx.shared import Inches
import os

def replace_text_in_doc(doc, key, value):
    """Replace text in paragraphs, tables, headers, and footers."""
    for paragraph in doc.paragraphs:
        replace_text_preserve_style(paragraph, key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_preserve_style(paragraph, key, str(value))

    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    replace_text_preserve_style(paragraph, key, str(value))
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    replace_text_preserve_style(paragraph, key, str(value))

def replace_text_preserve_style(paragraph, key, value):
    """Replace text in a paragraph while preserving formatting."""
    if key in paragraph.text:
        inline = paragraph.runs
        for i in range(len(inline)):
            if key in inline[i].text:
                text = inline[i].text.replace(key, value)
                inline[i].text = text

def insert_screenshot(doc, placeholder, image_path):
    """Replace a screenshot placeholder with an actual image."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, "")
            try:
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(6.0))
            except FileNotFoundError:
                print(f"Error: Screenshot not found at {image_path}")
            except Exception as e:
                print(f"An error occurred while inserting the image: {e}")

def process_document(template_path, output_path, data_dict, screenshot_path):
    """Process a single document with replacement data."""
    try:
        doc = Document(template_path)
        # Create a dictionary for text replacements, WITH underscores in placeholders
        replacements = {f"{{{{{k.replace(' ', '_').upper()}}}}}": v for k, v in data_dict.items()}

        # Replace placeholders in the entire document
        for key, value in replacements.items():
            replace_text_in_doc(doc, key, str(value))

        image_placeholder = '{{SCREENSHOT_OUTPUT}}'
        image_inserted = False

        # Check if a valid screenshot path was provided and the file exists
        if screenshot_path and os.path.exists(screenshot_path):
            for p in doc.paragraphs:
                if image_placeholder in p.text:
                    # Remove the placeholder text from the paragraph
                    p.text = p.text.replace(image_placeholder, '')
                    
                    # Add the picture to that same paragraph
                    run = p.add_run()
                    try:
                        run.add_picture(screenshot_path, width=Inches(6.0))
                        print(f"Successfully inserted screenshot into '{output_path}'")
                        image_inserted = True
                        break  # Exit the loop once the image has been placed
                    except Exception as e:
                        print(f"Error: Failed to add picture from '{screenshot_path}'. Reason: {e}")
                        # If insertion fails, write an error message in the document
                        run.text = f"[ERROR: Could not insert image. {e}]"
                        break

        doc.save(output_path)
        print(f"Document saved: {output_path}")

    except FileNotFoundError:
        print(f"Error: Template file not found at {template_path}")
    except Exception as e:
        print(f"An error occurred while processing {template_path}: {e}")

# def process_document(template_path, output_path, data_dict, screenshot_path):
#     """
#     Processes a Word document template, fills in placeholders, and saves the output.
#     This version correctly handles both text and image placeholders.
#     """
#     doc = Document(template_path)

#     # --- 1. Replace Text Placeholders in Paragraphs and Tables ---
#     # Combine all paragraphs from the main body and tables into one list
#     all_paragraphs = list(doc.paragraphs)
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 all_paragraphs.extend(cell.paragraphs)

#     # Replace text placeholders
#     for p in all_paragraphs:
#         for key, value in data_dict.items():
#             # This is a simple replacement. For text, this is often sufficient.
#             if f'{{{{{key}}}}}' in p.text:
#                 # Note: A more complex replacement would iterate through p.runs
#                 p.text = p.text.replace(f'{{{{{key}}}}}', str(value))

#     # --- 2. Replace Image Placeholder ---
#     image_placeholder = '{{SCREENSHOT_OUTPUT}}'
#     image_inserted = False

#     # Check if a valid screenshot path was provided and the file exists
#     if screenshot_path and os.path.exists(screenshot_path):
#         for p in doc.paragraphs:
#             if image_placeholder in p.text:
#                 # Remove the placeholder text from the paragraph
#                 p.text = p.text.replace(image_placeholder, '')
                
#                 # Add the picture to that same paragraph
#                 run = p.add_run()
#                 try:
#                     run.add_picture(screenshot_path, width=Inches(6.0))
#                     print(f"Successfully inserted screenshot into '{output_path}'")
#                     image_inserted = True
#                     break  # Exit the loop once the image has been placed
#                 except Exception as e:
#                     print(f"Error: Failed to add picture from '{screenshot_path}'. Reason: {e}")
#                     # If insertion fails, write an error message in the document
#                     run.text = f"[ERROR: Could not insert image. {e}]"
#                     break
    
#     # --- 3. Add Diagnostic Warnings if Image was Not Inserted ---
#     if not image_inserted and screenshot_path:
#         # This warning triggers if the path was valid but the placeholder wasn't found
#         if os.path.exists(screenshot_path):
#              print(f"Warning: Screenshot path was valid, but placeholder '{image_placeholder}' was not found in '{template_path}'.")

#     # --- 4. Save the Final Document ---
#     try:
#         doc.save(output_path)
#     except Exception as e:
#         print(f"Error saving document '{output_path}': {e}")

def main():
    """Main function to drive the document generation."""
    try:
        excel_file = "charm_data.xlsx"
        df = pd.read_excel(excel_file)
    except FileNotFoundError:
        print(f"Error: Excel file not found: {excel_file}")
        return
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return

    df['Description'] = df['Program Name']
    df['Test Plan Prepared By'] = df['Created By']
    df['Testing By'] = df['Test Plan Reviewed By']
    df['Test Result Prepared By'] = df['Created By']

    template_folder = "charm-docs-template"
    templates = {
        "Spec.docx": "Spec_Output.docx",
        "Test Plan.docx": "Test_Plan_Output.docx",
        "Test Results.docx": "Test_Results_Output.docx"
    }

    for index, row in df.iterrows():
        data_dict = row.to_dict()
        change_number = row['Change Number']
        screenshot_path = str(row.get('Screenshot Path', '')).strip().strip("'\"")

        for template_name, output_file in templates.items():
            template_path = os.path.join(template_folder, template_name)
            if os.path.exists(template_path):
                base_name, ext = os.path.splitext(output_file)
                unique_output = f"{base_name}_{change_number}{ext}"
                process_document(template_path, unique_output, data_dict, screenshot_path)
            else:
                print(f"Template file not found: {template_path}")

if __name__ == "__main__":
    main()
